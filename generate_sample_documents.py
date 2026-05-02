from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


BASE_DIR = Path(__file__).resolve().parent
REPO_DIR = BASE_DIR / "repository"


def ensure_dirs() -> None:
    for dept in ["finance", "procurement", "hr", "operations"]:
        (REPO_DIR / dept).mkdir(parents=True, exist_ok=True)


def generate_finance_documents():
    dept_dir = REPO_DIR / "finance"
    dept_dir.mkdir(parents=True, exist_ok=True)

    # -----------------------------
    # 1. DOCX
    # -----------------------------
    doc_path = dept_dir / "Finance_Report_Q1.docx"
    doc = Document()
    doc.add_heading("Finance Department – Q1 Report", level=0)
    doc.add_paragraph("This report summarises financial performance for Q1.")

    # Table
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    hdr[2].text = "Notes"

    rows = [
        ("Revenue", "12,450,000", "Strong performance"),
        ("Operating Income", "3,120,000", "Stable margins"),
        ("Net Profit", "2,480,000", "Healthy quarter"),
    ]
    for r in rows:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text = r

    # Chart
    chart_path = dept_dir / "Finance_Q1_Chart.png"
    plt.figure(figsize=(4, 3))
    plt.bar(["Revenue", "Op Income", "Net Profit"], [12.45, 3.12, 2.48])
    plt.ylabel("£ Millions")
    plt.title("Finance KPIs – Q1")
    plt.tight_layout()
    plt.savefig(chart_path)
    plt.close()

    doc.add_picture(str(chart_path), width=Inches(4.5))
    doc.save(doc_path)

    # -----------------------------
    # 2. XLSX
    # -----------------------------
    xlsx_path = dept_dir / "Finance_Report_Q1.xlsx"
    df = pd.DataFrame({
        "Metric": ["Revenue", "Operating Income", "Net Profit"],
        "Value": [12450000, 3120000, 2480000],
        "Notes": ["Strong", "Stable", "Healthy"]
    })

    with pd.ExcelWriter(xlsx_path, engine="xlsxwriter") as writer:
        df.to_excel(writer, sheet_name="Finance", index=False)
        workbook = writer.book
        worksheet = writer.sheets["Finance"]
        worksheet.insert_image("E2", str(chart_path))

    # -----------------------------
    # 3. PDF
    # -----------------------------
    pdf_path = dept_dir / "Finance_Report_Q1.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=A4)
    width, height = A4

    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, height - 50, "Finance Department – Q1 Report")

    y = height - 100
    for metric, value, note in rows:
        c.drawString(50, y, f"{metric}: {value} ({note})")
        y -= 20

    c.drawImage(str(chart_path), 50, y - 200, width=300, preserveAspectRatio=True)
    c.showPage()
    c.save()

    # -----------------------------
    # 4. PNG
    # -----------------------------
    png_path = dept_dir / "Finance_Report_Q1.png"
    img = Image.new("RGB", (900, 600), "white")
    draw = ImageDraw.Draw(img)
    draw.text((40, 40), "Finance Department – Q1 Report", fill="black")
    draw.text((40, 100), "Revenue: 12,450,000", fill="black")
    draw.text((40, 140), "Operating Income: 3,120,000", fill="black")
    draw.text((40, 180), "Net Profit: 2,480,000", fill="black")
    img.save(png_path)



def generate_procurement_documents():
    dept_dir = REPO_DIR / "procurement"
    dept_dir.mkdir(parents=True, exist_ok=True)

    # Chart
    chart_path = dept_dir / "Procurement_Q1_Chart.png"
    vendors = ["BlueTech", "RapidFix", "GreenLeaf"]
    spend = [320000, 48600, 89500]

    plt.figure(figsize=(4, 3))
    plt.bar(vendors, spend)
    plt.ylabel("Spend (£)")
    plt.title("Vendor Spend – Q1")
    plt.tight_layout()
    plt.savefig(chart_path)
    plt.close()

    # DOCX
    doc_path = dept_dir / "Procurement_Report_Q1.docx"
    doc = Document()
    doc.add_heading("Procurement – Q1 Vendor Spend", level=0)
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Vendor"
    hdr[1].text = "Spend (£)"
    for v, s in zip(vendors, spend):
        row = table.add_row().cells
        row[0].text = v
        row[1].text = f"{s:,}"
    doc.add_picture(str(chart_path), width=Inches(4.5))
    doc.save(doc_path)

    # XLSX
    xlsx_path = dept_dir / "Procurement_Report_Q1.xlsx"
    df = pd.DataFrame({"Vendor": vendors, "Spend": spend})
    with pd.ExcelWriter(xlsx_path, engine="xlsxwriter") as writer:
        df.to_excel(writer, sheet_name="Spend", index=False)
        writer.sheets["Spend"].insert_image("D2", str(chart_path))

    # PDF
    pdf_path = dept_dir / "Procurement_Report_Q1.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=A4)
    y = A4[1] - 50
    c.drawString(50, y, "Procurement – Q1 Vendor Spend")
    y -= 40
    for v, s in zip(vendors, spend):
        c.drawString(50, y, f"{v}: £{s:,}")
        y -= 20
    c.drawImage(str(chart_path), 50, y - 200, width=300)
    c.showPage()
    c.save()

    # PNG
    png_path = dept_dir / "Procurement_Report_Q1.png"
    img = Image.new("RGB", (900, 600), "white")
    draw = ImageDraw.Draw(img)
    draw.text((40, 40), "Procurement – Q1 Vendor Spend", fill="black")
    for i, (v, s) in enumerate(zip(vendors, spend)):
        draw.text((40, 100 + i * 40), f"{v}: £{s:,}", fill="black")
    img.save(png_path)


def generate_hr_documents():
    dept_dir = REPO_DIR / "hr"
    dept_dir.mkdir(parents=True, exist_ok=True)

    departments = ["Finance", "Procurement", "HR", "Operations", "Tech"]
    headcount = [48, 22, 18, 142, 82]

    chart_path = dept_dir / "HR_Headcount_Chart.png"
    plt.figure(figsize=(4, 3))
    plt.bar(departments, headcount)
    plt.title("Headcount by Department – Q1")
    plt.tight_layout()
    plt.savefig(chart_path)
    plt.close()

    # DOCX
    doc_path = dept_dir / "HR_Report_Q1.docx"
    doc = Document()
    doc.add_heading("HR – Q1 Payroll & Headcount", level=0)
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Department"
    hdr[1].text = "Headcount"
    for d, h in zip(departments, headcount):
        row = table.add_row().cells
        row[0].text = d
        row[1].text = str(h)
    doc.add_picture(str(chart_path), width=Inches(4.5))
    doc.save(doc_path)

    # XLSX
    xlsx_path = dept_dir / "HR_Report_Q1.xlsx"
    df = pd.DataFrame({"Department": departments, "Headcount": headcount})
    with pd.ExcelWriter(xlsx_path, engine="xlsxwriter") as writer:
        df.to_excel(writer, sheet_name="Headcount", index=False)
        writer.sheets["Headcount"].insert_image("D2", str(chart_path))

    # PDF
    pdf_path = dept_dir / "HR_Report_Q1.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=A4)
    y = A4[1] - 50
    c.drawString(50, y, "HR – Q1 Payroll & Headcount")
    y -= 40
    for d, h in zip(departments, headcount):
        c.drawString(50, y, f"{d}: {h}")
        y -= 20
    c.drawImage(str(chart_path), 50, y - 200, width=300)
    c.showPage()
    c.save()

    # PNG
    png_path = dept_dir / "HR_Report_Q1.png"
    img = Image.new("RGB", (900, 600), "white")
    draw = ImageDraw.Draw(img)
    draw.text((40, 40), "HR – Q1 Payroll & Headcount", fill="black")
    for i, (d, h) in enumerate(zip(departments, headcount)):
        draw.text((40, 100 + i * 40), f"{d}: {h}", fill="black")
    img.save(png_path)



def generate_operations_documents():
    dept_dir = REPO_DIR / "operations"
    dept_dir.mkdir(parents=True, exist_ok=True)

    services = ["Electrical", "Plumbing", "HVAC"]
    spend = [4860, 3200, 7800]

    chart_path = dept_dir / "Ops_Spend_Chart.png"
    plt.figure(figsize=(4, 3))
    plt.bar(services, spend)
    plt.title("Maintenance Spend – Q1")
    plt.tight_layout()
    plt.savefig(chart_path)
    plt.close()

    # DOCX
    doc_path = dept_dir / "Operations_Report_Q1.docx"
    doc = Document()
    doc.add_heading("Operations – Q1 Maintenance Spend", level=0)
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Service"
    hdr[1].text = "Spend (£)"
    for s, v in zip(services, spend):
        row = table.add_row().cells
        row[0].text = s
        row[1].text = f"{v:,}"
    doc.add_picture(str(chart_path), width=Inches(4.5))
    doc.save(doc_path)

    # XLSX
    xlsx_path = dept_dir / "Operations_Report_Q1.xlsx"
    df = pd.DataFrame({"Service": services, "Spend": spend})
    with pd.ExcelWriter(xlsx_path, engine="xlsxwriter") as writer:
        df.to_excel(writer, sheet_name="Maintenance", index=False)
        writer.sheets["Maintenance"].insert_image("D2", str(chart_path))

    # PDF
    pdf_path = dept_dir / "Operations_Report_Q1.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=A4)
    y = A4[1] - 50
    c.drawString(50, y, "Operations – Q1 Maintenance Spend")
    y -= 40
    for s, v in zip(services, spend):
        c.drawString(50, y, f"{s}: £{v:,}")
        y -= 20
    c.drawImage(str(chart_path), 50, y - 200, width=300)
    c.showPage()
    c.save()

    # PNG
    png_path = dept_dir / "Operations_Report_Q1.png"
    img = Image.new("RGB", (900, 600), "white")
    draw = ImageDraw.Draw(img)
    draw.text((40, 40), "Operations – Q1 Maintenance Spend", fill="black")
    for i, (s, v) in enumerate(zip(services, spend)):
        draw.text((40, 100 + i * 40), f"{s}: £{v:,}", fill="black")
    img.save(png_path)



if __name__ == "__main__":
    ensure_dirs()
    generate_finance_documents()
    generate_procurement_documents()
    generate_hr_documents()
    generate_operations_documents()
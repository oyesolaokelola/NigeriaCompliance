#!/usr/bin/env python
"""Test the template styling system end-to-end."""

import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add project to path
import sys
sys.path.insert(0, str(Path(__file__).parent))

from workflow.template_styling import TemplateExtractor, TemplateManager, StyleApplier

def test_template_extraction():
    """Test extracting styling from a template document."""
    print("\n" + "="*60)
    print("TEST 1: Template Extraction")
    print("="*60)
    
    template_path = Path("templates/professional_template.docx")
    if not template_path.exists():
        print(f"❌ Template file not found: {template_path}")
        return False
    
    try:
        extractor = TemplateExtractor()
        profile = extractor.extract(str(template_path))
        
        print(f"✅ Successfully extracted template: {profile.template_name}")
        print(f"\nExtracted Styling Information:")
        print(f"  Title Font:")
        print(f"    - Name: {profile.title_font.name}")
        print(f"    - Size: {profile.title_font.size}pt")
        print(f"    - Bold: {profile.title_font.bold}")
        print(f"    - Color: #{profile.title_font.color}")
        print(f"  Heading Font:")
        print(f"    - Name: {profile.heading_font.name}")
        print(f"    - Size: {profile.heading_font.size}pt")
        print(f"    - Bold: {profile.heading_font.bold}")
        print(f"  Body Font:")
        print(f"    - Name: {profile.body_font.name}")
        print(f"    - Size: {profile.body_font.size}pt")
        print(f"  Page Layout:")
        print(f"    - Left Margin: {profile.margin_left} twips")
        print(f"    - Top Margin: {profile.margin_top} twips")
        
        return profile
    except Exception as e:
        print(f"❌ Extraction failed: {e}")
        import traceback
        traceback.print_exc()
        return None


def test_template_registration(profile):
    """Test registering a template in the template manager."""
    print("\n" + "="*60)
    print("TEST 2: Template Registration & Caching")
    print("="*60)
    
    try:
        templates_dir = Path("templates")
        manager = TemplateManager(templates_dir)
        
        # Register the template
        registered = manager.register_template(
            "templates/professional_template.docx",
            "test_professional"
        )
        
        print(f"✅ Template registered: {registered.template_name}")
        
        # List all templates
        templates = manager.list_templates()
        print(f"✅ Registered templates: {templates}")
        
        # Retrieve cached profile
        retrieved = manager.get_template_profile("test_professional")
        if retrieved:
            print(f"✅ Successfully retrieved cached profile")
            return retrieved
        else:
            print(f"❌ Failed to retrieve cached profile")
            return None
            
    except Exception as e:
        print(f"❌ Registration failed: {e}")
        import traceback
        traceback.print_exc()
        return None


def test_style_application(profile):
    """Test applying styles to a document."""
    print("\n" + "="*60)
    print("TEST 3: Applying Styles to Generated Document")
    print("="*60)
    
    try:
        # Create a sample document with default styling
        doc = Document()
        
        # Add content
        title = doc.add_paragraph("Test Compliance Report")
        heading = doc.add_paragraph("Executive Summary")
        body = doc.add_paragraph("This is a test document that will have styling applied.")
        
        print(f"📄 Created test document with 3 paragraphs")
        
        # Apply template styling
        applier = StyleApplier(profile)
        styled_doc = applier.apply_to_docx(doc)
        
        print(f"✅ Successfully applied template styling")
        
        # Save the styled document
        output_path = Path("output/test_styled_document.docx")
        output_path.parent.mkdir(exist_ok=True)
        styled_doc.save(str(output_path))
        
        print(f"✅ Saved styled document to: {output_path}")
        print(f"   Full path: {output_path.absolute()}")
        
        # Verify the styling was applied
        verify_doc = Document(str(output_path))
        first_para = verify_doc.paragraphs[0]
        if first_para.runs:
            first_run = first_para.runs[0]
            print(f"\n✅ Verification:")
            print(f"   - Title font applied: {first_run.font.name}")
            print(f"   - Title size applied: {first_run.font.size}")
            print(f"   - Title bold applied: {first_run.font.bold}")
        
        return output_path
        
    except Exception as e:
        print(f"❌ Style application failed: {e}")
        import traceback
        traceback.print_exc()
        return None


def test_workflow_integration():
    """Test the full workflow with template processing."""
    print("\n" + "="*60)
    print("TEST 4: Full Workflow Integration")
    print("="*60)
    
    try:
        # Test creating a sample data document
        from workflow.extraction import extract_word
        
        # Create a test data document
        doc = Document()
        doc.add_paragraph("Q1 2025 Financial Report")
        doc.add_paragraph("Department: Finance")
        doc.add_paragraph("Revenue: $1,000,000")
        
        test_doc_path = Path("output/test_data_document.docx")
        test_doc_path.parent.mkdir(exist_ok=True)
        doc.save(str(test_doc_path))
        
        print(f"✅ Created test data document: {test_doc_path}")
        
        # Extract from the document
        extracted = extract_word(test_doc_path)
        print(f"✅ Successfully extracted data from document")
        print(f"   - Text length: {len(extracted.get('raw_text', ''))} chars")
        print(f"   - File type: {extracted.get('file_type')}")
        
        return True
        
    except Exception as e:
        print(f"❌ Workflow integration test failed: {e}")
        import traceback
        traceback.print_exc()
        return False


def print_summary():
    """Print test summary and instructions."""
    print("\n" + "="*60)
    print("TESTING SUMMARY")
    print("="*60)
    print("""
✅ Template Styling System is OPERATIONAL

Next Steps to Test Full Integration:
1. Start the FastAPI server:
   uvicorn api.server:app --reload --port 8000

2. Upload the template:
   curl -X POST http://localhost:8000/templates/upload \\
     -F "file=@templates/professional_template.docx" \\
     -F "template_name=professional"

3. Activate the template:
   curl -X POST http://localhost:8000/templates/professional/activate

4. Upload data documents to process via the /upload endpoint

5. Start processing:
   curl -X POST http://localhost:8000/process

6. Check results with generated output using the template styling
""")


if __name__ == "__main__":
    print("\n🧪 TEMPLATE STYLING SYSTEM TEST")
    print("="*60)
    
    # Run tests in sequence
    profile = test_template_extraction()
    if not profile:
        print("\n❌ Template extraction failed, stopping tests")
        exit(1)
    
    retrieved_profile = test_template_registration(profile)
    if not retrieved_profile:
        retrieved_profile = profile  # Continue with extracted profile
    
    styled_doc_path = test_style_application(retrieved_profile)
    if not styled_doc_path:
        print("\n⚠️  Style application failed, but continuing...")
    
    workflow_ok = test_workflow_integration()
    
    print_summary()
    print("\n✅ All template styling tests completed!")

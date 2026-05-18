# workflow/template_styling.py
"""
Template-based styling module for the NigeriaCompliance workflow.

This module provides functionality to extract styling information from template documents
and apply those styles to generated documents, ensuring consistent formatting across reports.
"""

import logging
import json
from pathlib import Path
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, field, asdict
from collections import Counter
import tempfile
import shutil

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
try:
    from docx.enum.section import WD_ORIENT
except ImportError:
    from docx.enum.text import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import pdfplumber

logger = logging.getLogger(__name__)


@dataclass
class FontStyle:
    """Extracted font styling information"""
    name: str = "Calibri"
    size: int = 11
    bold: bool = False
    italic: bool = False
    underline: bool = False
    color: str = "000000"  # RGB hex
    highlight_color: str = None  # RGB hex for highlight
    strike_through: bool = False
    subscript: bool = False
    superscript: bool = False
    

@dataclass
class ParagraphStyle:
    """Extracted paragraph styling information"""
    alignment: str = "left"  # left, center, right, justify
    line_spacing: float = 1.0
    space_before: int = 0
    space_after: int = 0
    indent_left: int = 0
    indent_right: int = 0
    indent_first_line: int = 0
    keep_with_next: bool = False
    page_break_before: bool = False
    widow_control: bool = True
    shading_color: str = None  # RGB hex for paragraph background


@dataclass
class TableStyle:
    """Extracted table styling information"""
    style_name: str = "Light Grid Accent 1"
    border_color: str = "000000"
    border_width: int = 1
    cell_padding: int = 0
    header_background_color: str = None
    header_font: FontStyle = None
    body_font: FontStyle = None
    banding: bool = True  # Alternating row colors
    banding_color: str = None
    
    def __post_init__(self):
        if self.header_font is None:
            self.header_font = FontStyle()
        if self.body_font is None:
            self.body_font = FontStyle()


@dataclass
class ListStyle:
    """Extracted list formatting information"""
    list_type: str = "bullet"  # bullet, numbered, multilevel
    bullet_char: str = "•"
    numbering_format: str = "1."  # 1., a., i., etc.
    indent_level: int = 0
    indent_hanging: int = 720  # Twips
    space_before: int = 0
    space_after: int = 0


@dataclass
class ImageStyle:
    """Extracted image styling information"""
    width: float = 0  # In inches
    height: float = 0  # In inches
    alignment: str = "left"
    wrap_text: bool = True
    position_x: float = 0  # Normalized 0-100
    position_y: float = 0  # Normalized 0-100
    

@dataclass
class SectionStyle:
    """Extracted section styling information"""
    heading_level: int = 1
    title_font: FontStyle = None
    body_font: FontStyle = None
    heading_font: FontStyle = None
    paragraph_style: ParagraphStyle = None
    margin_top: int = 720  # In twips (1/20th of a point)
    margin_bottom: int = 720
    

@dataclass
class TemplateProfile:
    """Complete styling profile extracted from a template document"""
    template_name: str
    template_path: str
    page_width: float = 8.5
    page_height: float = 11.0
    margin_left: int = 720
    margin_right: int = 720
    margin_top: int = 720
    margin_bottom: int = 720
    title_font: FontStyle = None
    heading_font: FontStyle = None
    body_font: FontStyle = None
    paragraph_style: ParagraphStyle = None
    section_styles: Dict[str, SectionStyle] = None
    table_style: TableStyle = None
    list_styles: Dict[str, ListStyle] = None
    image_styles: List[ImageStyle] = None
    color_scheme: Dict[str, str] = None
    header_content: str = ""
    footer_content: str = ""
    has_page_numbers: bool = False
    orientation: str = "portrait"
    logo_path: Optional[str] = None
    logo_style: ImageStyle = None
    implied_rules: List[str] = None
    template_insights: Dict[str, Any] = None
    custom_styles: Dict[str, Dict[str, Any]] = None  # All custom style definitions
    

    def __post_init__(self):
        if self.title_font is None:
            self.title_font = FontStyle()
        if self.heading_font is None:
            self.heading_font = FontStyle()
        if self.body_font is None:
            self.body_font = FontStyle()
        if self.paragraph_style is None:
            self.paragraph_style = ParagraphStyle()
        if self.section_styles is None:
            self.section_styles = {}
        if self.table_style is None:
            self.table_style = TableStyle()
        if self.list_styles is None:
            self.list_styles = {}
        if self.image_styles is None:
            self.image_styles = []
        if self.color_scheme is None:
            self.color_scheme = {}
        if self.logo_style is None:
            self.logo_style = ImageStyle()
        if self.implied_rules is None:
            self.implied_rules = []
        if self.template_insights is None:
            self.template_insights = {}
        if self.custom_styles is None:
            self.custom_styles = {}


class PDFConverter:
    """Utility for converting PDF to DOCX for enhanced template extraction"""
    
    @staticmethod
    def convert_pdf_to_docx(pdf_path: Path, output_dir: Path = None) -> Path:
        """
        Convert PDF to DOCX using pdf2docx library.
        
        Args:
            pdf_path: Path to the PDF file
            output_dir: Directory to save the converted DOCX (defaults to same as PDF)
            
        Returns:
            Path to the converted DOCX file
            
        Raises:
            ImportError: If pdf2docx is not installed
            Exception: If conversion fails
        """
        try:
            from pdf2docx import Converter
        except ImportError:
            logger.warning("pdf2docx library not available, using fallback conversion")
            return PDFConverter.convert_pdf_to_docx_fallback(pdf_path, output_dir)
        
        if output_dir is None:
            output_dir = pdf_path.parent
        
        docx_path = output_dir / f"{pdf_path.stem}.docx"
        
        try:
            logger.info(f"Converting PDF to DOCX: {pdf_path}")
            cv = Converter(str(pdf_path))
            cv.convert(str(docx_path))
            cv.close()
            logger.info(f"Successfully converted PDF to DOCX: {docx_path}")
            return docx_path
        except Exception as e:
            logger.error(f"Error converting PDF to DOCX: {e}")
            # Fall back to basic conversion if pdf2docx fails
            logger.info("Falling back to basic PDF conversion")
            return PDFConverter.convert_pdf_to_docx_fallback(pdf_path, output_dir)
    
    @staticmethod
    def convert_pdf_to_docx_fallback(pdf_path: Path, output_dir: Path = None) -> Path:
        """
        Fallback method to convert PDF to DOCX using pdfplumber.
        This is less accurate but works without pdf2docx.
        
        Args:
            pdf_path: Path to the PDF file
            output_dir: Directory to save the converted DOCX (defaults to same as PDF)
            
        Returns:
            Path to the converted DOCX file
        """
        if output_dir is None:
            output_dir = pdf_path.parent
        
        docx_path = output_dir / f"{pdf_path.stem}.docx"
        
        try:
            logger.info(f"Converting PDF to DOCX using fallback method: {pdf_path}")
            doc = Document()
            
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        doc.add_paragraph(text)
                    
                    # Extract tables
                    tables = page.extract_tables()
                    for table_data in tables:
                        if table_data:
                            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                            for i, row in enumerate(table_data):
                                for j, cell in enumerate(row):
                                    table.rows[i].cells[j].text = str(cell) if cell else ""
            
            doc.save(docx_path)
            logger.info(f"Successfully converted PDF to DOCX using fallback: {docx_path}")
            return docx_path
        except Exception as e:
            logger.error(f"Error converting PDF to DOCX using fallback: {e}")
            raise


class TemplateExtractor:
    """Extracts styling information from template documents"""
    
    def __init__(self):
        self.supported_formats = ['.docx', '.pdf']
        self.converter = PDFConverter()
    
    def extract_from_docx(self, template_path: Path) -> TemplateProfile:
        """Extract styling from a DOCX template with enhanced accuracy"""
        try:
            doc = Document(template_path)
            profile = TemplateProfile(
                template_name=template_path.stem,
                template_path=str(template_path)
            )
            
            # Extract page layout
            section = doc.sections[0]
            profile.margin_left = section.left_margin.twips
            profile.margin_right = section.right_margin.twips
            profile.margin_top = section.top_margin.twips
            profile.margin_bottom = section.bottom_margin.twips
            profile.orientation = "landscape" if section.orientation == WD_ORIENT.LANDSCAPE else "portrait"
            
            # Extract all custom style definitions from the document
            profile.custom_styles = self._extract_all_styles(doc)
            
            # Collect font samples with enhanced attributes
            font_samples = []
            for para in doc.paragraphs[:50]:
                if not para.text.strip():
                    continue
                for run in para.runs:
                    font = run.font
                    font_samples.append(FontStyle(
                        name=font.name or "Calibri",
                        size=int(font.size.pt) if font.size else 11,
                        bold=bool(font.bold),
                        italic=bool(font.italic),
                        underline=bool(font.underline),
                        color=self._extract_font_color(font),
                        highlight_color=self._extract_highlight_color(run),
                        strike_through=bool(font.strike),
                        subscript=bool(font.subscript),
                        superscript=bool(font.superscript)
                    ))
            
            # Use multi-factor ranking for font detection
            ranked = self._rank_font_styles_enhanced(font_samples)
            profile.title_font = ranked.get("title", FontStyle(name="Calibri", size=28, bold=True))
            profile.heading_font = ranked.get("heading", FontStyle(name="Calibri", size=14, bold=True))
            profile.body_font = ranked.get("body", FontStyle(name="Calibri", size=11))
            profile.color_scheme = {
                "primary": profile.title_font.color,
                "secondary": profile.heading_font.color,
                "body": profile.body_font.color,
            }
            
            # Extract paragraph styles from multiple paragraphs
            paragraph_styles = []
            for para in doc.paragraphs[:20]:
                if para.text.strip():
                    pPr = para._element.pPr
                    if pPr is not None:
                        paragraph_styles.append(self._extract_paragraph_style(pPr))
            
            # Use the most common paragraph style
            if paragraph_styles:
                profile.paragraph_style = self._select_most_common_paragraph_style(paragraph_styles)
            
            # Extract table styles with detailed information
            if doc.tables:
                profile.table_style = self._extract_table_style(doc.tables[0])
            
            # Extract list styles
            profile.list_styles = self._extract_list_styles(doc)
            
            # Extract image styles
            profile.image_styles = self._extract_image_styles(doc)
            
            # Extract header/footer text and styles
            if doc.sections[0].header.paragraphs:
                header_text = " ".join([p.text for p in doc.sections[0].header.paragraphs if p.text.strip()])
                profile.header_content = header_text[:150]
            if doc.sections[0].footer.paragraphs:
                footer_text = " ".join([p.text for p in doc.sections[0].footer.paragraphs if p.text.strip()])
                profile.footer_content = footer_text[:150]
            profile.has_page_numbers = "page" in profile.footer_content.lower()
            
            # Extract logo or first image from the template with positioning
            profile.logo_path, profile.logo_style = self._extract_docx_logo_with_style(doc, template_path.parent, profile.template_name)
            
            # Extract implicit template rules from structure and content
            heading_texts = [para.text.strip() for para in doc.paragraphs if para.style and para.style.name and "heading" in para.style.name.lower()]
            body_texts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            table_headers = []
            for table in doc.tables:
                if table.rows:
                    row = table.rows[0]
                    table_headers.extend([cell.text.strip() for cell in row.cells if cell.text.strip()])
            profile.implied_rules = self._infer_template_rules(body_texts, heading_texts, table_headers, bool(profile.logo_path))
            profile.template_insights = {
                "headings": heading_texts[:10],
                "table_headers": table_headers[:10],
                "has_logo": bool(profile.logo_path),
                "total_paragraphs": len(doc.paragraphs),
                "total_tables": len(doc.tables),
                "total_images": len(profile.image_styles),
            }
            
            logger.info(f"Successfully extracted enhanced template profile from {template_path}")
            return profile
        except Exception as e:
            logger.error(f"Error extracting DOCX template: {e}")
            raise
    
    def extract_from_pdf(self, template_path: Path) -> TemplateProfile:
        """Extract styling information from a PDF template"""
        try:
            profile = TemplateProfile(
                template_name=template_path.stem,
                template_path=str(template_path)
            )
            
            with pdfplumber.open(template_path) as pdf:
                if pdf.pages:
                    page = pdf.pages[0]
                    profile.page_width = page.width / 72  # Convert to inches
                    profile.page_height = page.height / 72
                    
                    font_samples = []
                    for obj in page.chars[:200]:
                        font_name = obj.get('fontname') or obj.get('name') or 'Helvetica'
                        font_size = int(obj.get('size', 11))
                        font_samples.append(FontStyle(
                            name=font_name,
                            size=font_size,
                            bold=False,
                            italic=False,
                            color="000000"
                        ))
                    
                    ranked = self._rank_font_styles(font_samples)
                    profile.title_font = ranked.get("title", FontStyle(name="Helvetica", size=28, bold=True))
                    profile.heading_font = ranked.get("heading", FontStyle(name="Helvetica", size=14, bold=True))
                    profile.body_font = ranked.get("body", FontStyle(name="Helvetica", size=11))
                    profile.color_scheme = {
                        "primary": profile.title_font.color,
                        "secondary": profile.heading_font.color,
                        "body": profile.body_font.color,
                    }
                    
                    profile.logo_path = self._extract_pdf_logo(page, template_path.parent, profile.template_name)
                    page_text = page.extract_text() or ""
                    text_snippets = [line.strip() for line in page_text.splitlines() if line.strip()]
                    profile.implied_rules = self._infer_template_rules(text_snippets, [], [], bool(profile.logo_path))
                    profile.template_insights = {
                        "text_snippets": text_snippets[:10],
                        "has_logo": bool(profile.logo_path),
                    }
            
            logger.info(f"Successfully extracted template profile from PDF {template_path}")
            return profile
        except Exception as e:
            logger.error(f"Error extracting PDF template: {e}")
            raise
    
    def extract(self, template_path: str, convert_pdf: bool = True) -> TemplateProfile:
        """
        Extract template profile from file (auto-detect format).
        
        Args:
            template_path: Path to the template file
            convert_pdf: If True, convert PDF to DOCX before extraction for enhanced accuracy
            
        Returns:
            TemplateProfile with extracted styling information
        """
        path = Path(template_path)
        if not path.exists():
            raise FileNotFoundError(f"Template file not found: {template_path}")
        
        suffix = path.suffix.lower()
        if suffix == '.docx':
            return self.extract_from_docx(path)
        elif suffix == '.pdf':
            if convert_pdf:
                # Convert PDF to DOCX for enhanced extraction
                try:
                    docx_path = self.converter.convert_pdf_to_docx(path)
                    logger.info(f"Using converted DOCX for enhanced extraction: {docx_path}")
                    profile = self.extract_from_docx(docx_path)
                    # Update template path to point to the original PDF
                    profile.template_path = str(path)
                    return profile
                except ImportError:
                    logger.warning("pdf2docx not available, using fallback conversion")
                    docx_path = self.converter.convert_pdf_to_docx_fallback(path)
                    profile = self.extract_from_docx(docx_path)
                    profile.template_path = str(path)
                    return profile
                except Exception as e:
                    logger.warning(f"PDF conversion failed, using basic PDF extraction: {e}")
                    return self.extract_from_pdf(path)
            else:
                return self.extract_from_pdf(path)
        else:
            raise ValueError(f"Unsupported template format: {suffix}")
    
    @staticmethod
    def _extract_font_color(font) -> str:
        """Extract RGB color from font"""
        try:
            if font.color and font.color.rgb:
                return font.color.rgb.string.lstrip('#')
            return "000000"
        except:
            return "000000"
    
    @staticmethod
    def _extract_highlight_color(run) -> str:
        """Extract highlight color from text run"""
        try:
            if run.element.highlight:
                color = run.element.highlight.val
                # Convert highlight color to RGB hex
                highlight_colors = {
                    'yellow': 'FFFF00',
                    'brightGreen': '00FF00',
                    'turquoise': '00FFFF',
                    'pink': 'FF00FF',
                    'blue': '0000FF',
                    'red': 'FF0000',
                    'darkBlue': '00008B',
                    'teal': '008080',
                    'gray': '808080',
                    'darkGray': 'A9A9A9',
                    'lightGray': 'D3D3D3',
                    'black': '000000',
                }
                return highlight_colors.get(color, None)
            return None
        except:
            return None
    
    @staticmethod
    def _extract_all_styles(doc: Document) -> Dict[str, Dict[str, Any]]:
        """Extract all custom style definitions from the document"""
        styles = {}
        try:
            for style in doc.styles:
                if style.type == 1:  # Paragraph style
                    # Convert RGBColor to hex string
                    color_hex = '000000'
                    if style.font.color and style.font.color.rgb:
                        try:
                            color_hex = style.font.color.rgb.string.lstrip('#')
                        except:
                            color_hex = '000000'
                    
                    styles[style.name] = {
                        'type': 'paragraph',
                        'font': {
                            'name': style.font.name,
                            'size': style.font.size.pt if style.font.size else 11,
                            'bold': style.font.bold,
                            'italic': style.font.italic,
                            'color': color_hex
                        },
                        'paragraph': {
                            'alignment': str(style.paragraph_format.alignment) if style.paragraph_format.alignment else 'left',
                            'line_spacing': style.paragraph_format.line_spacing,
                            'space_before': style.paragraph_format.space_before.pt if style.paragraph_format.space_before else 0,
                            'space_after': style.paragraph_format.space_after.pt if style.paragraph_format.space_after else 0,
                        }
                    }
                elif style.type == 2:  # Character style
                    # Convert RGBColor to hex string
                    color_hex = '000000'
                    if style.font.color and style.font.color.rgb:
                        try:
                            color_hex = style.font.color.rgb.string.lstrip('#')
                        except:
                            color_hex = '000000'
                    
                    styles[style.name] = {
                        'type': 'character',
                        'font': {
                            'name': style.font.name,
                            'size': style.font.size.pt if style.font.size else 11,
                            'bold': style.font.bold,
                            'italic': style.font.italic,
                            'color': color_hex
                        }
                    }
        except Exception as e:
            logger.warning(f"Error extracting styles: {e}")
        return styles
    
    def _rank_font_styles_enhanced(self, font_styles: List[FontStyle]) -> Dict[str, FontStyle]:
        """
        Rank font styles using multi-factor analysis (size, bold, italic, color).
        """
        if not font_styles:
            return {}

        # Group by size first
        size_groups: Dict[int, List[FontStyle]] = {}
        for font in font_styles:
            size_groups.setdefault(font.size, []).append(font)

        unique_sizes = sorted(size_groups.keys(), reverse=True)
        if not unique_sizes:
            return {}

        # Title: Largest size, usually bold
        title_candidates = size_groups.get(unique_sizes[0], [])
        title_font = self._select_most_likely_title(title_candidates)

        # Heading: Second largest or largest non-bold
        heading_candidates = size_groups.get(unique_sizes[1] if len(unique_sizes) > 1 else unique_sizes[0], [])
        heading_font = self._select_most_likely_heading(heading_candidates)

        # Body: Most common smaller size
        body_size = unique_sizes[-1]
        body_candidates = size_groups.get(body_size, [])
        body_font = self._select_most_common_style(body_candidates)

        return {
            "title": title_font,
            "heading": heading_font,
            "body": body_font,
        }
    
    def _select_most_likely_title(self, styles: List[FontStyle]) -> FontStyle:
        """Select the most likely title font from candidates"""
        if not styles:
            return FontStyle(name="Calibri", size=28, bold=True)
        
        # Prefer bold, larger fonts
        bold_styles = [s for s in styles if s.bold]
        if bold_styles:
            return self._select_most_common_style(bold_styles)
        return self._select_most_common_style(styles)
    
    def _select_most_likely_heading(self, styles: List[FontStyle]) -> FontStyle:
        """Select the most likely heading font from candidates"""
        if not styles:
            return FontStyle(name="Calibri", size=14, bold=True)
        
        # Prefer bold fonts for headings
        bold_styles = [s for s in styles if s.bold]
        if bold_styles:
            return self._select_most_common_style(bold_styles)
        return self._select_most_common_style(styles)
    
    def _select_most_common_paragraph_style(self, styles: List[ParagraphStyle]) -> ParagraphStyle:
        """Select the most common paragraph style from a list"""
        if not styles:
            return ParagraphStyle()
        
        counts = Counter(
            (s.alignment, s.line_spacing, s.space_before, s.space_after)
            for s in styles
        )
        best_style_key = max(counts.items(), key=lambda item: item[1])[0]
        for style in styles:
            if (style.alignment, style.line_spacing, style.space_before, style.space_after) == best_style_key:
                return style
        return styles[0]
    
    def _extract_table_style(self, table) -> TableStyle:
        """Extract detailed table styling information"""
        table_style = TableStyle()
        
        try:
            # Get table style name
            if table.style:
                table_style.style_name = table.style.name
            
            # Extract border information from first cell
            if table.rows and table.rows[0].cells:
                first_cell = table.rows[0].cells[0]
                tcPr = first_cell._element.tcPr
                if tcPr is not None:
                    tcBorders = tcPr.find(qn('w:tcBorders'))
                    if tcBorders is not None:
                        # Extract border color and width
                        top_border = tcBorders.find(qn('w:top'))
                        if top_border is not None:
                            table_style.border_color = top_border.get(qn('w:val'), '000000')
                            table_style.border_width = int(top_border.get(qn('w:sz'), '4')) / 8  # Convert to points
            
            # Extract header font from first row
            if table.rows:
                header_row = table.rows[0]
                if header_row.cells and header_row.cells[0].paragraphs:
                    first_para = header_row.cells[0].paragraphs[0]
                    if first_para.runs:
                        run = first_para.runs[0]
                        table_style.header_font = FontStyle(
                            name=run.font.name or "Calibri",
                            size=int(run.font.size.pt) if run.font.size else 11,
                            bold=bool(run.font.bold),
                            italic=bool(run.font.italic),
                            color=self._extract_font_color(run.font)
                        )
            
            # Extract body font from second row if available
            if len(table.rows) > 1:
                body_row = table.rows[1]
                if body_row.cells and body_row.cells[0].paragraphs:
                    first_para = body_row.cells[0].paragraphs[0]
                    if first_para.runs:
                        run = first_para.runs[0]
                        table_style.body_font = FontStyle(
                            name=run.font.name or "Calibri",
                            size=int(run.font.size.pt) if run.font.size else 11,
                            bold=bool(run.font.bold),
                            italic=bool(run.font.italic),
                            color=self._extract_font_color(run.font)
                        )
        
        except Exception as e:
            logger.warning(f"Error extracting table style: {e}")
        
        return table_style
    
    def _extract_list_styles(self, doc: Document) -> Dict[str, ListStyle]:
        """Extract list formatting styles from the document"""
        list_styles = {}
        
        try:
            for para in doc.paragraphs:
                if para.style and para.style.name:
                    style_name = para.style.name.lower()
                    if 'list' in style_name or 'bullet' in style_name:
                        pPr = para._element.pPr
                        if pPr is not None:
                            numPr = pPr.find(qn('w:numPr'))
                            if numPr is not None:
                                list_style = ListStyle()
                                
                                # Extract numbering format
                                ilvl = numPr.find(qn('w:ilvl'))
                                if ilvl is not None:
                                    list_style.indent_level = int(ilvl.get(qn('w:val'), '0'))
                                
                                numId = numPr.find(qn('w:numId'))
                                if numId is not None:
                                    # Try to determine list type from numbering
                                    list_style.list_type = "numbered" if numId else "bullet"
                                
                                list_styles[para.style.name] = list_style
        except Exception as e:
            logger.warning(f"Error extracting list styles: {e}")
        
        return list_styles
    
    def _extract_image_styles(self, doc: Document) -> List[ImageStyle]:
        """Extract image styling information from the document"""
        image_styles = []
        
        try:
            for rel in doc.part.rels.values():
                if 'image' in rel.reltype:
                    try:
                        image_part = rel.target_part
                        # Get image dimensions if available
                        width = 0
                        height = 0
                        
                        # Try to get dimensions from the relationship
                        if hasattr(image_part, '_part'):
                            try:
                                from PIL import Image
                                import io
                                img = Image.open(io.BytesIO(image_part.blob))
                                width, height = img.size
                                # Convert pixels to inches (assuming 96 DPI)
                                width = width / 96
                                height = height / 96
                            except:
                                pass
                        
                        image_style = ImageStyle(
                            width=width,
                            height=height,
                            alignment="left",
                            wrap_text=True
                        )
                        image_styles.append(image_style)
                    except Exception:
                        continue
        except Exception as e:
            logger.warning(f"Error extracting image styles: {e}")
        
        return image_styles
    
    def _extract_docx_logo_with_style(self, doc: Document, template_dir: Path, template_name: str) -> tuple:
        """Extract logo with positioning information"""
        try:
            for rel in doc.part.rels.values():
                if 'image' in rel.reltype:
                    try:
                        image_part = rel.target_part
                        image_bytes = image_part.blob
                        ext = image_part.content_type.split('/')[-1]
                        logo_name = f"{template_name}_logo.{ext}"
                        logo_path = template_dir / logo_name
                        with open(logo_path, "wb") as fh:
                            fh.write(image_bytes)
                        
                        # Extract image style
                        logo_style = ImageStyle()
                        try:
                            from PIL import Image
                            import io
                            img = Image.open(io.BytesIO(image_bytes))
                            width, height = img.size
                            logo_style.width = width / 96  # Convert to inches
                            logo_style.height = height / 96
                        except:
                            pass
                        
                        return str(logo_path), logo_style
                    except Exception:
                        continue
        except Exception:
            pass
        return None, ImageStyle()
    
    @staticmethod
    def _extract_paragraph_style(pPr) -> ParagraphStyle:
        """Extract enhanced paragraph properties from OOXML element"""
        style = ParagraphStyle()
        
        # Alignment
        jc = pPr.find(qn('w:jc'))
        if jc is not None:
            align_map = {'left': 'left', 'center': 'center', 'right': 'right', 'both': 'justify'}
            style.alignment = align_map.get(jc.get(qn('w:val')), 'left')
        
        # Line spacing
        spacing = pPr.find(qn('w:spacing'))
        if spacing is not None:
            line = spacing.get(qn('w:line'))
            if line:
                style.line_spacing = int(line) / 240
        
        # Space before/after
        if spacing is not None:
            before = spacing.get(qn('w:before'))
            after = spacing.get(qn('w:after'))
            if before:
                style.space_before = int(before)
            if after:
                style.space_after = int(after)
        
        # Indentation
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            left = ind.get(qn('w:left'))
            right = ind.get(qn('w:right'))
            first_line = ind.get(qn('w:firstLine'))
            if left:
                style.indent_left = int(left)
            if right:
                style.indent_right = int(right)
            if first_line:
                style.indent_first_line = int(first_line)
        
        # Page break control
        keep_next = pPr.find(qn('w:keepNext'))
        if keep_next is not None:
            style.keep_with_next = keep_next.get(qn('w:val')) == '1'
        
        page_break_before = pPr.find(qn('w:pageBreakBefore'))
        if page_break_before is not None:
            style.page_break_before = page_break_before.get(qn('w:val')) == '1'
        
        widow_control = pPr.find(qn('w:widowControl'))
        if widow_control is not None:
            style.widow_control = widow_control.get(qn('w:val')) != '0'
        
        # Shading
        shd = pPr.find(qn('w:shd'))
        if shd is not None:
            fill = shd.get(qn('w:fill'))
            if fill:
                style.shading_color = fill
        
        return style

    def _infer_template_rules(self, body_texts: List[str], heading_texts: List[str], table_headers: List[str], has_logo: bool) -> List[str]:
        """Infer implicit reporting and compliance rules from a template document."""
        rules: List[str] = []
        lower_text = " ".join(body_texts + heading_texts + table_headers).lower()
        keywords = {
            "revenue": "Include revenue or sales figures in the report when present.",
            "payroll": "Capture payroll and labor cost metrics if they appear in the template.",
            "vat": "Match VAT or tax-related calculations according to the template structure.",
            "invoice": "Treat invoice line items and totals as core financial data.",
            "vendor": "Consider procurement and vendor risk metrics when vendor sections appear.",
            "risk": "Generate a risk analysis section when the template includes risk or compliance headings.",
            "compliance": "Preserve compliance-related sections and summary blocks from the template.",
            "net profit": "Use net profit or net amount as a key financial summary when available.",
            "operational": "Include operations metrics if operational sections are present.",
            "headcount": "Capture HR or headcount details if the template includes workforce information.",
        }

        for keyword, rule_text in keywords.items():
            if keyword in lower_text and rule_text not in rules:
                rules.append(rule_text)

        if heading_texts:
            rules.append("Honor the template section and heading structure when generating the report.")
        if table_headers:
            rules.append("Preserve table column semantics based on template table headers.")
        if has_logo:
            rules.insert(0, "Retain template branding and logo placement in generated outputs.")

        if not rules:
            rules.append("No explicit template rules were inferred; generate output based on the template style and structure.")

        return rules

    def _rank_font_styles(self, font_styles: List[FontStyle]) -> Dict[str, FontStyle]:
        """Rank font styles dynamically based on template usage."""
        if not font_styles:
            return {}

        size_groups: Dict[int, List[FontStyle]] = {}
        for font in font_styles:
            size_groups.setdefault(font.size, []).append(font)

        unique_sizes = sorted(size_groups.keys(), reverse=True)
        if not unique_sizes:
            return {}

        title_size = unique_sizes[0]
        heading_size = unique_sizes[1] if len(unique_sizes) > 1 else title_size
        body_size = unique_sizes[-1]

        title_font = self._select_most_common_style(size_groups[title_size])
        heading_font = self._select_most_common_style(size_groups[heading_size])
        body_font = self._select_most_common_style(size_groups[body_size])

        return {
            "title": title_font,
            "heading": heading_font,
            "body": body_font,
        }

    @staticmethod
    def _select_most_common_style(styles: List[FontStyle]) -> FontStyle:
        if not styles:
            return FontStyle()

        counts = Counter(
            (style.name, style.size, style.bold, style.italic, style.color)
            for style in styles
        )
        best_style_key = max(counts.items(), key=lambda item: (item[1], item[0][1]))[0]
        for style in styles:
            if (style.name, style.size, style.bold, style.italic, style.color) == best_style_key:
                return style
        return styles[0]

    @staticmethod
    def _extract_docx_logo(doc: Document, template_dir: Path, template_name: str) -> Optional[str]:
        for rel in doc.part.rels.values():
            if 'image' in rel.reltype:
                try:
                    image_part = rel.target_part
                    image_bytes = image_part.blob
                    ext = image_part.content_type.split('/')[-1]
                    logo_name = f"{template_name}_logo.{ext}"
                    logo_path = template_dir / logo_name
                    with open(logo_path, "wb") as fh:
                        fh.write(image_bytes)
                    return str(logo_path)
                except Exception:
                    continue
        return None

    @staticmethod
    def _extract_pdf_logo(page, template_dir: Path, template_name: str) -> Optional[str]:
        images = getattr(page, 'images', None) or []
        if not images:
            return None

        image_obj = images[0]
        try:
            image_data = page.extract_image(image_obj.get('object_id'))
            if not image_data:
                return None
            image_bytes = image_data.get('image')
            ext = image_data.get('ext', 'png')
            logo_name = f"{template_name}_logo.{ext}"
            logo_path = template_dir / logo_name
            with open(logo_path, "wb") as fh:
                fh.write(image_bytes)
            return str(logo_path)
        except Exception:
            return None


class StyleApplier:
    """Applies extracted template styles to generated documents"""
    
    def __init__(self, template_profile: TemplateProfile):
        self.profile = template_profile
    
    def apply_html_content(self, doc: Document, html_content: str, layout_elements: List[Dict[str, Any]] = None) -> Document:
        """
        Apply HTML content from multimodal parser to DOCX document.
        
        Args:
            doc: The DOCX document to modify
            html_content: HTML content containing tables and structured elements
            layout_elements: Optional list of layout elements with bounding boxes
            
        Returns:
            Modified DOCX document
        """
        try:
            # Parse HTML content
            try:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(html_content, 'html.parser')
            except ImportError:
                logger.warning("BeautifulSoup not available, cannot parse HTML content")
                return doc
            
            # Process each div with layout information
            if layout_elements:
                for element in layout_elements:
                    self._apply_layout_element(doc, element)
            
            # Process tables from HTML
            for table in soup.find_all('table'):
                self._add_html_table_to_docx(doc, table)
            
            # Process text content
            for div in soup.find_all('div'):
                # Skip if it's a table container
                if div.find('table'):
                    continue
                
                # Get the category from data-label if available
                category = div.get('data-label', 'Text')
                content = div.get_text(strip=True)
                
                if content:
                    self._add_content_by_category(doc, content, category)
            
            logger.info("Successfully applied HTML content to document")
            return doc
            
        except Exception as e:
            logger.error(f"Error applying HTML content: {e}")
            return doc
    
    def _add_html_table_to_docx(self, doc: Document, html_table):
        """
        Convert an HTML table with colspan/rowspan to a DOCX table.
        
        Args:
            doc: The DOCX document
            html_table: BeautifulSoup table element
        """
        try:
            rows = html_table.find_all('tr')
            if not rows:
                return
            
            # Determine grid dimensions accounting for colspan/rowspan
            max_cols = 0
            for row in rows:
                col_count = 0
                for cell in row.find_all(['td', 'th']):
                    col_count += int(cell.get('colspan', 1))
                max_cols = max(max_cols, col_count)
            
            if max_cols == 0:
                return
            
            # Create DOCX table
            table = doc.add_table(rows=len(rows), cols=max_cols)
            table.style = self.profile.table_style
            
            # Track which cells are already filled due to rowspan
            cell_tracker = [[False for _ in range(max_cols)] for _ in range(len(rows))]
            
            # Fill the table
            for row_idx, row in enumerate(rows):
                docx_row = table.rows[row_idx]
                col_idx = 0
                
                for cell in row.find_all(['td', 'th']):
                    # Find next available column
                    while col_idx < max_cols and cell_tracker[row_idx][col_idx]:
                        col_idx += 1
                    
                    if col_idx >= max_cols:
                        break
                    
                    colspan = int(cell.get('colspan', 1))
                    rowspan = int(cell.get('rowspan', 1))
                    cell_text = cell.get_text(strip=True)
                    
                    # Set cell text
                    docx_cell = docx_row.cells[col_idx]
                    docx_cell.text = cell_text
                    
                    # Mark cells affected by rowspan
                    for r in range(row_idx, min(row_idx + rowspan, len(rows))):
                        for c in range(col_idx, min(col_idx + colspan, max_cols)):
                            if r < len(cell_tracker) and c < len(cell_tracker[r]):
                                cell_tracker[r][c] = True
                    
                    # Apply colspan in DOCX (merge cells horizontally)
                    if colspan > 1 and col_idx + colspan <= max_cols:
                        merged_cell = docx_row.cells[col_idx]
                        for c in range(col_idx + 1, min(col_idx + colspan, max_cols)):
                            try:
                                merged_cell.merge(docx_row.cells[c])
                            except:
                                pass
                    
                    col_idx += colspan
            
        except Exception as e:
            logger.error(f"Error converting HTML table to DOCX: {e}")
    
    def _apply_layout_element(self, doc: Document, element: Dict[str, Any]):
        """
        Apply a layout element with bounding box information.
        
        Args:
            doc: The DOCX document
            element: Dictionary with bbox, category, and content
        """
        category = element.get('category', 'Text')
        content = element.get('content', '').strip()
        
        if content:
            self._add_content_by_category(doc, content, category)
    
    def _add_content_by_category(self, doc: Document, content: str, category: str):
        """
        Add content to document based on its category.
        
        Args:
            doc: The DOCX document
            content: The text content
            category: The element category (Title, Section-header, Text, Table, etc.)
        """
        if category == 'Title':
            para = doc.add_paragraph(content)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._apply_font_style(para, self.profile.title_font)
        elif category == 'Section-header':
            para = doc.add_paragraph(content)
            self._apply_font_style(para, self.profile.heading_font)
        elif category == 'Text':
            para = doc.add_paragraph(content)
            self._apply_font_style(para, self.profile.body_font)
            self._apply_paragraph_style(para, self.profile.paragraph_style)
        elif category == 'Caption':
            para = doc.add_paragraph(content)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._apply_font_style(para, FontStyle(name=self.profile.body_font.name, 
                                                   size=self.profile.body_font.size - 1,
                                                   italic=True))
        else:
            # Default to body text for other categories
            para = doc.add_paragraph(content)
            self._apply_font_style(para, self.profile.body_font)
    
    def apply_to_docx(self, doc: Document) -> Document:
        """Apply template styling to a DOCX document"""
        try:
            # Apply margins
            for section in doc.sections:
                section.left_margin = self.profile.margin_left
                section.right_margin = self.profile.margin_right
                section.top_margin = self.profile.margin_top
                section.bottom_margin = self.profile.margin_bottom
                
                # Apply header/footer if template has them
                if self.profile.header_content:
                    header = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
                    header.text = self.profile.header_content
                
                if self.profile.footer_content or self.profile.has_page_numbers:
                    footer = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
                    if self.profile.has_page_numbers:
                        footer.text = self.profile.footer_content or "Page "
            
            # Apply styles to paragraphs
            for i, para in enumerate(doc.paragraphs):
                style_name = para.style.name.lower() if para.style and para.style.name else ""
                if "title" in style_name or "heading 1" in style_name or (i == 0 and not style_name):
                    self._apply_font_style(para, self.profile.title_font)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif "heading" in style_name:
                    self._apply_font_style(para, self.profile.heading_font)
                else:
                    self._apply_font_style(para, self.profile.body_font)
                    self._apply_paragraph_style(para, self.profile.paragraph_style)
            
            # Apply table styles with enhanced styling
            for table in doc.tables:
                self._apply_table_style(table, self.profile.table_style)
            
            logger.info("Successfully applied template styles to document")
            return doc
            
        except Exception as e:
            logger.error(f"Error applying template styles: {e}")
            return doc
    
    def apply_to_template(self, content_data: Dict[str, Any], output_path: Path) -> Document:
        """
        Generate document by modifying the template in-place for accurate mirroring.
        
        This approach preserves exact layout, styling, and structure from the template.
        
        Args:
            content_data: Dictionary containing content to insert:
                - title: Document title
                - sections: List of section dictionaries with 'heading' and 'content'
                - tables: List of table data
                - metadata: Additional metadata like date, author, etc.
            output_path: Path where the generated document will be saved
            
        Returns:
            Generated Document object
        """
        try:
            # Load the template document
            template_doc = Document(self.profile.template_path)
            
            # Replace title if present
            if content_data.get('title'):
                for para in template_doc.paragraphs[:5]:  # Check first 5 paragraphs for title
                    if para.text.strip() and any(keyword in para.text.lower() for keyword in ['title', 'report', 'document']):
                        para.text = content_data['title']
                        break
            
            # Replace metadata placeholders
            metadata = content_data.get('metadata', {})
            for para in template_doc.paragraphs:
                for key, value in metadata.items():
                    placeholder = f"{{{key}}}"
                    if placeholder in para.text:
                        para.text = para.text.replace(placeholder, str(value))
            
            # Add sections content
            sections = content_data.get('sections', [])
            if sections:
                # Find where to insert content (after title/intro)
                insert_index = 0
                for i, para in enumerate(template_doc.paragraphs):
                    if para.text.strip():
                        insert_index = i + 1
                        break
                
                for section in sections:
                    # Add heading
                    if section.get('heading'):
                        heading_para = template_doc.add_paragraph(section['heading'])
                        heading_para.style = 'Heading 1'
                        self._apply_font_style(heading_para, self.profile.heading_font)
                    
                    # Add content
                    if section.get('content'):
                        content_para = template_doc.add_paragraph(section['content'])
                        self._apply_font_style(content_para, self.profile.body_font)
                        self._apply_paragraph_style(content_para, self.profile.paragraph_style)
            
            # Add tables
            tables_data = content_data.get('tables', [])
            if tables_data:
                for table_data in tables_data:
                    self._add_table_to_template(template_doc, table_data)
            
            # Save the document
            template_doc.save(output_path)
            logger.info(f"Successfully generated document from template: {output_path}")
            
            return template_doc
            
        except Exception as e:
            logger.error(f"Error generating document from template: {e}")
            raise
    
    def _add_table_to_template(self, doc: Document, table_data: Dict[str, Any]):
        """
        Add a table to the template document with proper styling.
        
        Args:
            doc: The template document
            table_data: Dictionary containing table data:
                - headers: List of column headers
                - rows: List of row data (list of lists)
        """
        try:
            headers = table_data.get('headers', [])
            rows = table_data.get('rows', [])
            
            if not headers or not rows:
                return
            
            # Create table
            table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
            table.style = self.profile.table_style.style_name
            
            # Apply enhanced table styling
            self._apply_table_style(table, self.profile.table_style)
            
            # Add headers
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = str(header)
            
            # Add data rows
            for row_idx, row_data in enumerate(rows):
                for col_idx, cell_data in enumerate(row_data):
                    table.rows[row_idx + 1].cells[col_idx].text = str(cell_data)
            
        except Exception as e:
            logger.error(f"Error adding table to template: {e}")
    
    @staticmethod
    def _apply_font_style(para, font_style: FontStyle):
        """Apply enhanced font styling to a paragraph"""
        for run in para.runs:
            run.font.name = font_style.name
            run.font.size = Pt(font_style.size)
            run.font.bold = font_style.bold
            run.font.italic = font_style.italic
            run.font.underline = font_style.underline
            run.font.strike = font_style.strike_through
            run.font.subscript = font_style.subscript
            run.font.superscript = font_style.superscript
            
            # Set color
            try:
                if font_style.color and len(font_style.color) == 6:
                    rgb = bytes.fromhex(font_style.color)
                    if len(rgb) == 3:
                        run.font.color.rgb = RGBColor(*rgb)
            except:
                pass
            
            # Set highlight color if available
            if font_style.highlight_color:
                try:
                    from docx.oxml.shared import OxmlElement
                    rPr = run._element.get_or_add_rPr()
                    highlight = OxmlElement('w:highlight')
                    highlight.set(qn('w:val'), font_style.highlight_color)
                    rPr.append(highlight)
                except:
                    pass
    
    @staticmethod
    def _apply_paragraph_style(para, para_style: ParagraphStyle):
        """Apply enhanced paragraph styling"""
        pPr = para._element.get_or_add_pPr()
        
        # Set spacing
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:line'), str(int(para_style.line_spacing * 240)))
        spacing.set(qn('w:before'), str(para_style.space_before))
        spacing.set(qn('w:after'), str(para_style.space_after))
        pPr.append(spacing)
        
        # Set indentation
        if para_style.indent_left > 0 or para_style.indent_right > 0 or para_style.indent_first_line != 0:
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), str(para_style.indent_left))
            ind.set(qn('w:right'), str(para_style.indent_right))
            ind.set(qn('w:firstLine'), str(para_style.indent_first_line))
            pPr.append(ind)
        
        # Set alignment
        jc = OxmlElement('w:jc')
        align_map = {'left': 'left', 'center': 'center', 'right': 'right', 'justify': 'both'}
        jc.set(qn('w:val'), align_map.get(para_style.alignment, 'left'))
        pPr.append(jc)
        
        # Set page break control
        if para_style.keep_with_next:
            keep_next = OxmlElement('w:keepNext')
            keep_next.set(qn('w:val'), '1')
            pPr.append(keep_next)
        
        if para_style.page_break_before:
            page_break = OxmlElement('w:pageBreakBefore')
            page_break.set(qn('w:val'), '1')
            pPr.append(page_break)
        
        if not para_style.widow_control:
            widow_control = OxmlElement('w:widowControl')
            widow_control.set(qn('w:val'), '0')
            pPr.append(widow_control)
        
        # Set shading if available
        if para_style.shading_color:
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), para_style.shading_color)
            pPr.append(shd)
    
    @staticmethod
    def _apply_table_style(table, table_style: TableStyle):
        """Apply enhanced table styling with custom borders, shading, and fonts"""
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Apply table style name
        tblStyle = OxmlElement('w:tblStyle')
        tblStyle.set(qn('w:val'), table_style.style_name)
        tblPr.append(tblStyle)
        
        # Apply custom borders if specified
        if table_style.border_color or table_style.border_width:
            tblBorders = OxmlElement('w:tblBorders')
            
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), str(table_style.border_width * 8))  # Convert to eighths of a point
                border.set(qn('w:color'), table_style.border_color)
                tblBorders.append(border)
            
            tblPr.append(tblBorders)
        
        # Apply header row styling
        if table.rows:
            header_row = table.rows[0]
            for cell in header_row.cells:
                # Apply header background color
                if table_style.header_background_color:
                    tcPr = cell._element.tcPr
                    if tcPr is None:
                        tcPr = OxmlElement('w:tcPr')
                        cell._element.append(tcPr)
                    
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), table_style.header_background_color)
                    tcPr.append(shd)
                
                # Apply header font
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = table_style.header_font.name
                        run.font.size = Pt(table_style.header_font.size)
                        run.font.bold = table_style.header_font.bold
                        run.font.italic = table_style.header_font.italic
                        try:
                            if table_style.header_font.color and len(table_style.header_font.color) == 6:
                                rgb = bytes.fromhex(table_style.header_font.color)
                                if len(rgb) == 3:
                                    run.font.color.rgb = RGBColor(*rgb)
                        except:
                            pass
        
        # Apply body row styling
        for i, row in enumerate(table.rows[1:], start=1):
            # Apply banding if enabled
            if table_style.banding and table_style.banding_color and i % 2 == 0:
                for cell in row.cells:
                    tcPr = cell._element.tcPr
                    if tcPr is None:
                        tcPr = OxmlElement('w:tcPr')
                        cell._element.append(tcPr)
                    
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), table_style.banding_color)
                    tcPr.append(shd)
            
            # Apply body font
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = table_style.body_font.name
                        run.font.size = Pt(table_style.body_font.size)
                        run.font.bold = table_style.body_font.bold
                        run.font.italic = table_style.body_font.italic
                        try:
                            if table_style.body_font.color and len(table_style.body_font.color) == 6:
                                rgb = bytes.fromhex(table_style.body_font.color)
                                if len(rgb) == 3:
                                    run.font.color.rgb = RGBColor(*rgb)
                        except:
                            pass


class TemplateManager:
    """Manages template storage and retrieval"""
    
    def __init__(self, template_dir: Path = None):
        if template_dir is None:
            template_dir = Path(__file__).resolve().parent.parent / "templates"
        self.template_dir = template_dir
        self.template_dir.mkdir(exist_ok=True, parents=True)
        self.profiles_cache_file = self.template_dir / "profiles.json"
    
    def register_template(self, template_path: str, profile_name: str = None) -> TemplateProfile:
        """Extract and register a template"""
        extractor = TemplateExtractor()
        profile = extractor.extract(template_path)
        
        if profile_name:
            profile.template_name = profile_name
        
        # Save profile to cache
        self._save_profile(profile)
        logger.info(f"Registered template: {profile.template_name}")
        return profile
    
    def get_template_profile(self, template_name: str) -> Optional[TemplateProfile]:
        """Retrieve a cached template profile"""
        profiles = self._load_profiles()
        for profile_dict in profiles.get("templates", []):
            if profile_dict.get("template_name") == template_name:
                return self._dict_to_profile(profile_dict)
        return None
    
    def list_templates(self) -> List[str]:
        """List all registered templates"""
        profiles = self._load_profiles()
        return [p.get("template_name") for p in profiles.get("templates", [])]
    
    def _save_profile(self, profile: TemplateProfile):
        """Save profile to cache file"""
        profiles = self._load_profiles()
        profile_dict = asdict(profile)
        
        # Convert nested dataclasses
        if profile.title_font:
            profile_dict['title_font'] = asdict(profile.title_font)
        if profile.heading_font:
            profile_dict['heading_font'] = asdict(profile.heading_font)
        if profile.body_font:
            profile_dict['body_font'] = asdict(profile.body_font)
        if profile.paragraph_style:
            profile_dict['paragraph_style'] = asdict(profile.paragraph_style)
        
        # Update or add
        existing = None
        for i, p in enumerate(profiles.get("templates", [])):
            if p.get("template_name") == profile.template_name:
                existing = i
                break
        
        if existing is not None:
            profiles["templates"][existing] = profile_dict
        else:
            profiles.setdefault("templates", []).append(profile_dict)
        
        with open(self.profiles_cache_file, 'w', encoding='utf-8') as f:
            json.dump(profiles, f, indent=2)
    
    def get_default_profile(self) -> TemplateProfile:
        """Return a safe default styling profile when no dynamic template is available."""
        default = TemplateProfile(
            template_name="default",
            template_path=str(self.template_dir / "default_profile"),
            margin_left=1440,
            margin_right=1440,
            margin_top=1440,
            margin_bottom=1440,
            title_font=FontStyle(name="Calibri", size=28, bold=True, color="000000"),
            heading_font=FontStyle(name="Calibri", size=16, bold=True, color="000000"),
            body_font=FontStyle(name="Calibri", size=11, bold=False, italic=False, color="000000"),
            paragraph_style=ParagraphStyle(alignment="left", line_spacing=1.15, space_before=0, space_after=120),
            color_scheme={"primary": "000000", "secondary": "333333", "body": "000000"},
            logo_path=None,
            orientation="portrait"
        )
        return default

    def get_template_or_default(self, template_name: Optional[str]) -> TemplateProfile:
        if template_name:
            profile = self.get_template_profile(template_name)
            if profile:
                return profile
        return self.get_default_profile()

    def _load_profiles(self) -> dict:
        """Load all cached profiles"""
        if self.profiles_cache_file.exists():
            try:
                with open(self.profiles_cache_file, 'r') as f:
                    return json.load(f)
            except:
                return {"templates": []}
        return {"templates": []}
    
    @staticmethod
    def _dict_to_profile(data: dict) -> TemplateProfile:
        """Convert dict to TemplateProfile"""
        # Handle nested dataclasses
        for key in ['title_font', 'heading_font', 'body_font', 'paragraph_style']:
            if key in data and isinstance(data[key], dict):
                if key == 'paragraph_style':
                    data[key] = ParagraphStyle(**data[key])
                else:
                    data[key] = FontStyle(**data[key])
        
        return TemplateProfile(**data)

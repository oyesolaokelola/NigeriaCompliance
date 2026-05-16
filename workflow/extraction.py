# workflow/extraction.py
from pathlib import Path
from typing import Dict, Any, Optional, List

import pandas as pd
from docx import Document
import pytesseract
from PIL import Image
import logging

logger = logging.getLogger(__name__)


def _safe_ocr_image(img: Image.Image) -> str:
    try:
        return pytesseract.image_to_string(img)
    except Exception:
        return ""


def extract_word(path: Path) -> Dict[str, Any]:
    doc = Document(path)
    text_lines = [p.text for p in doc.paragraphs]
    text = "\n".join(text_lines)
    tables: List[List[List[str]]] = []

    for t in doc.tables:
        table_data = []
        for row in t.rows:
            table_data.append([cell.text for cell in row.cells])
        tables.append(table_data)

    return {
        "raw_text": text,
        "raw_tables": tables,
        "source_path": str(path),
        "file_type": "docx",
    }


def extract_excel(path: Path) -> Dict[str, Any]:
    xls = pd.ExcelFile(path)
    tables = {}
    for sheet in xls.sheet_names:
        df = xls.parse(sheet)
        tables[sheet] = df.to_dict(orient="records")

    return {
        "raw_text": "",
        "raw_tables": tables,
        "source_path": str(path),
        "file_type": "xlsx",
    }


def extract_csv(path: Path) -> Dict[str, Any]:
    df = pd.read_csv(path)
    return {
        "raw_text": "",
        "raw_tables": {"csv": df.to_dict(orient="records")},
        "source_path": str(path),
        "file_type": "csv",
    }


def extract_text_file(path: Path) -> Dict[str, Any]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    return {
        "raw_text": text,
        "raw_tables": [],
        "source_path": str(path),
        "file_type": "txt",
    }


def extract_pdf(path: Path, use_multimodal_parser: bool = False, parser_config: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    """
    Extract content from a PDF file.
    
    Args:
        path: Path to the PDF file
        use_multimodal_parser: If True, use the multimodal LLM parser for accurate layout extraction
        parser_config: Configuration dict for multimodal parser (model_provider, model, etc.)
    
    Returns:
        Dictionary containing raw_text, raw_tables, source_path, file_type, and optionally
        html_content and layout_elements if using multimodal parser
    """
    # Use multimodal parser if requested
    if use_multimodal_parser:
        try:
            from .multimodal_parser import MultimodalParser
            
            config = parser_config or {}
            parser = MultimodalParser(
                model_provider=config.get("model_provider", "openai"),
                model=config.get("model", "gpt-4o"),
                reasoning_effort=config.get("reasoning_effort", "low"),
                merge_table=config.get("merge_table", True),
                create_html=config.get("create_html", True),
                additional_instructions=config.get("additional_instructions"),
                api_key=config.get("api_key")
            )
            
            result = parser.parse(str(path))
            
            # Extract tables from HTML if available
            tables = _extract_tables_from_html(result.clean_markdown) if result.clean_markdown else []
            
            return {
                "raw_text": result.clean_markdown,
                "raw_tables": tables,
                "source_path": str(path),
                "file_type": "pdf",
                "html_content": result.html,
                "layout_elements": result.layout_elements,
                "usage": {
                    "input_tokens": result.usage.input_tokens,
                    "output_tokens": result.usage.output_tokens,
                    "estimated_cost_usd": result.usage.estimated_cost_usd
                }
            }
        except ImportError:
            logger.warning("Multimodal parser requested but not available, falling back to pdfplumber")
        except Exception as e:
            logger.error(f"Multimodal parser failed: {e}, falling back to pdfplumber")
    
    # Fallback to traditional extraction
    text_chunks = []
    tables: List[List[List[str]]] = []
    text_chars = 0

    # Primary method: use pdfplumber (best for text + tables + page images).
    # If pdfplumber import or usage fails (e.g., incompatible native dependencies),
    # fall back to PyPDF2 for text-only extraction.
    try:
        import pdfplumber

        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                text_chunks.append(t)
                text_chars += len(t)
                try:
                    page_tables = page.extract_tables() or []
                    for tbl in page_tables:
                        tables.append(tbl)
                except Exception:
                    pass
    except Exception:
        # Fallback: try PyPDF2 for basic text extraction (no table extraction).
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(str(path))
            for page in reader.pages:
                try:
                    t = page.extract_text() or ""
                except Exception:
                    t = ""
                text_chunks.append(t)
                text_chars += len(t)
        except Exception:
            # Last-resort: leave text empty (OCR fallback below may also be skipped
            # if pdfplumber is not available to render page images).
            pass

    text = "\n".join(text_chunks)

    # If almost no text was extracted, try OCR via pdfplumber page images (only
    # possible when pdfplumber is available). If pdfplumber is not available,
    # we skip OCR fallback to avoid adding heavy native deps.
    if text_chars < 50:
        try:
            import pdfplumber

            images_text = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    pil_img = page.to_image(resolution=300).original
                    ocr_text = _safe_ocr_image(pil_img)
                    images_text.append(ocr_text)
            text = "\n".join(images_text)
        except Exception:
            pass

    return {
        "raw_text": text,
        "raw_tables": tables,
        "source_path": str(path),
        "file_type": "pdf",
    }


def _extract_tables_from_html(html_content: str) -> List[List[List[str]]]:
    """
    Extract table data from HTML content containing HTML tables.
    
    This is a simple implementation that parses HTML tables and converts
    them to the nested list format expected by the extraction pipeline.
    """
    try:
        from bs4 import BeautifulSoup
        
        soup = BeautifulSoup(html_content, 'html.parser')
        tables_data = []
        
        for table in soup.find_all('table'):
            table_data = []
            for row in table.find_all('tr'):
                row_data = []
                for cell in row.find_all(['td', 'th']):
                    # Handle colspan and rowspan by repeating cells
                    colspan = int(cell.get('colspan', 1))
                    rowspan = int(cell.get('rowspan', 1))
                    cell_text = cell.get_text(strip=True)
                    
                    # Simple handling: repeat cell text for colspan
                    for _ in range(colspan):
                        row_data.append(cell_text)
                
                if row_data:
                    table_data.append(row_data)
            
            if table_data:
                tables_data.append(table_data)
        
        return tables_data
    except ImportError:
        logger.warning("BeautifulSoup not available, cannot extract tables from HTML")
        return []
    except Exception as e:
        logger.error(f"Error extracting tables from HTML: {e}")
        return []


def extract_image(path: Path) -> Dict[str, Any]:
    img = Image.open(path)
    text = _safe_ocr_image(img)
    return {
        "raw_text": text,
        "raw_tables": [],
        "source_path": str(path),
        "file_type": "image",
    }


def extract_record(file_info: Dict[str, Any], use_multimodal_parser: bool = False, parser_config: Optional[Dict[str, Any]] = None) -> Optional[Dict[str, Any]]:
    """
    Extract content from a file record.
    
    Args:
        file_info: Dictionary containing file path and metadata
        use_multimodal_parser: If True, use multimodal parser for PDF files
        parser_config: Configuration dict for multimodal parser
    
    Returns:
        Dictionary containing extracted content or None if unsupported format
    """
    path = file_info["path"]
    suffix = path.suffix.lower()

    if suffix == ".docx":
        base = extract_word(path)
    elif suffix in [".xlsx", ".xls"]:
        base = extract_excel(path)
    elif suffix == ".pdf":
        base = extract_pdf(path, use_multimodal_parser=use_multimodal_parser, parser_config=parser_config)
    elif suffix == ".csv":
        base = extract_csv(path)
    elif suffix == ".txt":
        base = extract_text_file(path)
    elif suffix in [".png", ".jpg", ".jpeg"]:
        base = extract_image(path)
    else:
        return None

    # Minimal envelope; GenAI will classify department/period
    base.update(
        {
            "department": None,
            "period": None,
            "metrics": {},
            "notes": [],
        }
    )
    return base
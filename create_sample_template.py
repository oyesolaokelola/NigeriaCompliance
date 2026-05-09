#!/usr/bin/env python
"""Create a sample professional template for testing template styling."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

# Create a professional template document
doc = Document()

# Set up margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Add title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("Professional Report Template")
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue

# Add subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run("Compliance Analysis Report")
subtitle_run.font.size = Pt(14)
subtitle_run.font.italic = True
subtitle_run.font.color.rgb = RGBColor(100, 100, 100)  # Gray

# Add some spacing
doc.add_paragraph()

# Add section heading
heading1 = doc.add_paragraph()
heading_run = heading1.add_run("Executive Summary")
heading_run.font.size = Pt(16)
heading_run.font.bold = True
heading_run.font.color.rgb = RGBColor(0, 51, 102)

# Add body text
body = doc.add_paragraph(
    "This template demonstrates professional styling that will be applied to generated reports. "
    "It includes custom fonts, colors, spacing, and formatting conventions."
)
body_run = body.runs[0]
body_run.font.size = Pt(12)
body_run.font.name = "Calibri"

# Add another section
heading2 = doc.add_paragraph()
heading2_run = heading2.add_run("Key Findings")
heading2_run.font.size = Pt(16)
heading2_run.font.bold = True
heading2_run.font.color.rgb = RGBColor(0, 51, 102)

# Add bullet points
for item in ["Finding 1: Sample item", "Finding 2: Another item", "Finding 3: Third item"]:
    bullet = doc.add_paragraph(item, style='List Bullet')
    bullet.paragraph_format.left_indent = Inches(0.25)

# Add a table
table = doc.add_table(rows=3, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Metric'
hdr_cells[1].text = 'Value'
table.rows[1].cells[0].text = 'Status'
table.rows[1].cells[1].text = 'Compliant'
table.rows[2].cells[0].text = 'Issues'
table.rows[2].cells[1].text = 'None'

# Add footer
doc.add_paragraph()
footer = doc.add_paragraph("Generated using Professional Report Template")
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.runs[0]
footer_run.font.size = Pt(9)
footer_run.font.italic = True
footer_run.font.color.rgb = RGBColor(150, 150, 150)

# Save template
templates_dir = Path("templates")
templates_dir.mkdir(exist_ok=True)
template_path = templates_dir / "professional_template.docx"
doc.save(str(template_path))
print(f"✅ Sample template created: {template_path}")
print(f"Template location: {template_path.absolute()}")
